from __future__ import annotations

from datetime import date
from pathlib import Path
import sys

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = ROOT / "docs" / "Color-Tools-User-Manual.docx"
SCREENSHOT_DIR = ROOT / "docs" / "截图"


def set_run_font(run, size: float = 11, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold


def set_doc_default_font(document: Document) -> None:
    styles = document.styles
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def format_paragraph(paragraph, align=None, first_line_indent_cm: float | None = 0.74) -> None:
    paragraph.paragraph_format.line_spacing = 1.5
    if first_line_indent_cm is not None:
        paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    if align is not None:
        paragraph.alignment = align


def add_paragraph(document: Document, text: str, *, size: float = 11, bold: bool = False,
                  align=None, first_line_indent_cm: float | None = 0.74) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    format_paragraph(paragraph, align=align, first_line_indent_cm=first_line_indent_cm)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, bold=True)
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run, size=11)
        paragraph.paragraph_format.line_spacing = 1.4


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)

    document.add_paragraph("")


def add_toc(document: Document) -> None:
    """插入 Word 自动目录域（需在 Word 中更新域后生成页码与条目）。"""
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    r_element = run._r

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    # \o "1-2" 只取标题 1-2 级；\h 超链接；\u 使用大纲级别
    instr_text.text = r'TOC \o "1-2" \h \u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    r_element.append(fld_char_begin)
    r_element.append(instr_text)
    r_element.append(fld_char_separate)
    r_element.append(fld_char_end)


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not image_path.exists():
        return

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=10.5, bold=False)
    caption_paragraph.paragraph_format.space_after = Pt(6)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    set_doc_default_font(doc)

    core = doc.core_properties
    core.title = "色彩工坊使用说明书"
    core.subject = "Color Tools 用户使用说明"
    core.author = "Cursor AI Assistant"
    core.comments = "基于项目实际功能生成"

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(80)
    run = cover.add_run("色彩工坊")
    set_run_font(run, size=24, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("使用说明书")
    set_run_font(run, size=22, bold=True)

    target = doc.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target.paragraph_format.space_before = Pt(24)
    run = target.add_run("适用对象：终端用户 / 测试人员 / 项目交付文档")
    set_run_font(run, size=12)

    doc_date = doc.add_paragraph()
    doc_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = doc_date.add_run(f"生成日期：{date.today().isoformat()}")
    set_run_font(run, size=12)

    doc.add_page_break()

    # 自动目录页：标题 + TOC 域
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("目录")
    set_run_font(run, size=18, bold=True)

    add_toc(doc)

    doc.add_page_break()

    add_heading(doc, "一、引言", 1)
    add_heading(doc, "（一）编写目的", 2)
    add_paragraph(
        doc,
        "本说明书用于指导用户快速了解并使用“色彩工坊”桌面应用，使用户能够在无需额外培训的前提下完成颜色拾取、颜色转换、配色生成、传统色浏览、无障碍检查、调色板管理等常见操作。本文档既可用于项目交付，也可作为内部测试和日常使用参考。",
    )
    add_heading(doc, "（二）背景", 2)
    add_paragraph(
        doc,
        "“色彩工坊”是当前项目对应的综合色彩工具桌面应用，采用 Electron、React 与 Vite 构建，主要面向设计、前端开发、UI 设计、印刷配色、内容制作等场景。应用将图片取色、屏幕取色、格式转换、渐变参考、传统色查询、色盲模拟、对比度检查和调色板管理整合在统一界面中，帮助用户提升用色效率与协作便利性。",
    )

    add_heading(doc, "二、用途", 1)
    add_heading(doc, "（一）功能", 2)
    add_paragraph(doc, "本项目提供的主要功能如下：", first_line_indent_cm=None)
    add_bullets(
        doc,
        [
            "颜色拾取：支持图片取色与屏幕取色，输出 HEX、RGB 等结果，并保留最近取色记录。",
            "颜色计算与处理：支持彩图转黑白、单色转灰度、颜色混合、中间色生成、相反色计算、对比度检查和色盲模拟。",
            "颜色查询与参考：支持 HEX / RGB / HSL / HSV / CMYK 等常见色值转换，提供高级颜色选择器、渐变色集合、Web 安全色、RGB 与 CMYK 参考表。",
            "传统色彩浏览：支持中国传统色与日本传统色查看，并可直接复制色值。",
            "调色板管理：支持新建色板、重命名、添加/删除颜色、导入/导出 JSON 文件以及导出全部色板。",
            "通用能力：支持多语言界面、主题切换、快捷键跳转、反馈入口及更新检查等辅助功能。",
        ],
    )
    add_heading(doc, "（二）性能特点", 2)
    add_paragraph(
        doc,
        "本项目的大部分颜色计算和色值转换在本地实时完成，适合频繁试色、快速查询与日常设计辅助使用。图片处理速度会随着图片尺寸、像素数量和设备性能变化而有所差异。桌面端支持全局快捷键调起常用模块，可进一步提高操作效率。界面同时支持浅色、深色及跟随系统模式，并提供多语言能力，便于不同使用环境下的访问。",
    )

    add_heading(doc, "三、运行环境", 1)
    add_heading(doc, "（一）硬件设备建议", 2)
    add_paragraph(doc, "为获得较好的使用体验，建议采用如下硬件配置：", first_line_indent_cm=None)
    add_table(
        doc,
        ["项目", "建议配置"],
        [
            ["CPU", "双核及以上处理器"],
            ["内存", "4GB 及以上，建议 8GB"],
            ["存储空间", "预留 500MB 及以上可用空间"],
            ["显示分辨率", "建议 1366 × 768 及以上；应用最小窗口为 1024 × 768"],
        ],
    )
    add_heading(doc, "（二）支持软件", 2)
    add_paragraph(
        doc,
        "当前项目的桌面端打包目标以 Windows 平台为主，推荐在 Windows 10 或 Windows 11 环境下使用。对于安装包用户，无需额外安装开发依赖即可直接运行。对于源码运行或测试场景，需要预先安装 Node.js 与 npm，然后通过项目命令启动开发环境。",
    )
    add_table(
        doc,
        ["场景", "说明"],
        [
            ["安装包使用", "双击安装程序或便携版可执行文件即可运行"],
            ["源码开发", "执行 npm install 安装依赖，再按需运行开发命令"],
            ["前端开发端口", "Vite 默认端口为 5177"],
            ["桌面端打包", "可生成 Windows 安装包、ZIP 压缩包及解包目录"],
        ],
    )

    add_heading(doc, "四、操作说明", 1)
    add_heading(doc, "（一）PC 端应用安装与启动", 2)
    add_paragraph(
        doc,
        "如使用安装包版本，用户可双击安装程序，根据安装向导完成安装，并通过桌面快捷方式或开始菜单启动“色彩工坊”。如使用便携版，则可在解压后直接运行主程序。首次进入应用后，建议先确认界面语言、主题模式及常用功能位置。",
    )
    add_paragraph(doc, "如当前为项目内部测试或源码运行场景，可按以下方式启动：", first_line_indent_cm=None)
    add_table(
        doc,
        ["步骤", "命令 / 操作"],
        [
            ["1", "执行 npm install 安装项目依赖"],
            ["2", "执行 npm run electron:dev 启动桌面开发模式"],
            ["3", "如仅需前端页面调试，可执行 npm run dev"],
            ["4", "如需构建静态资源，可执行 npm run build"],
            ["5", "如需生成桌面安装包，可执行 npm run electron:build"],
        ],
    )
    add_heading(doc, "（二）首页与导航", 2)
    add_paragraph(
        doc,
        "应用首页用于展示主要能力入口。当前项目按照模块分为首页、颜色拾取、计算处理、颜色查询、传统色彩等主导航区域。用户可通过左侧或顶部导航进入对应页面，页面之间切换时会保留统一的界面风格与操作反馈。",
    )
    add_paragraph(
        doc,
        "顶部区域通常提供语言切换、主题切换、用户入口或辅助功能入口。用户在使用过程中可根据需要切换界面语言，并在浅色、深色或跟随系统模式之间选择合适的显示风格。",
    )
    add_image(doc, SCREENSHOT_DIR / "01-首页总览.png", "图 1 首页总览")
    add_heading(doc, "（三）颜色拾取", 2)
    add_paragraph(
        doc,
        "1. 图片取色：进入“颜色拾取”中的“图片取色”页面，上传待分析图片后，在预览区域点击任意像素位置即可获取对应颜色。系统会显示当前颜色的 HEX 与 RGB 结果，用户可一键复制所需色值，并查看最近取色历史。",
    )
    add_image(doc, SCREENSHOT_DIR / "02-图片取色.png", "图 2 图片取色")
    add_paragraph(
        doc,
        "2. 屏幕取色：进入“颜色拾取”中的“屏幕取色”页面，点击开始取色后，在屏幕上选择目标颜色即可获得结果。用户也可以使用快捷键直接跳转到该模块。若当前运行环境不支持相关取色能力，系统会给出不可用提示。",
    )
    add_image(doc, SCREENSHOT_DIR / "03-屏幕取色.png", "图 3 屏幕取色")
    add_heading(doc, "（四）颜色计算与处理", 2)
    add_paragraph(doc, "“计算处理”模块主要用于颜色运算和图像辅助处理，包含以下能力：", first_line_indent_cm=None)
    add_bullets(
        doc,
        [
            "彩图转黑白：上传图片后，系统会根据亮度算法生成黑白结果，并支持导出 PNG 文件。",
            "单色转灰度：输入指定颜色后，系统可自动输出其灰度结果。",
            "颜色混合：分别选择两种颜色并设置比例，系统会计算混合后的色值。",
            "中间色生成：输入起始颜色和结束颜色后，可生成一组连续过渡色。",
            "对比度检查：输入前景色与背景色后，可查看对比度数值及是否满足 WCAG AA / AAA 参考要求。",
            "相反色计算：输入一个颜色后，系统可给出其互补或相反色结果。",
            "色盲模拟：支持对单色或图片进行色盲显示效果模拟，辅助验证颜色可识别性。",
        ],
    )
    add_image(doc, SCREENSHOT_DIR / "04-彩图转黑白.png", "图 4 彩图转黑白")
    add_image(doc, SCREENSHOT_DIR / "05-颜色混合.png", "图 5 颜色混合")
    add_image(doc, SCREENSHOT_DIR / "06-中间色生成.png", "图 6 中间色生成")
    add_image(doc, SCREENSHOT_DIR / "07-对比度检查.png", "图 7 对比度检查")
    add_image(doc, SCREENSHOT_DIR / "08-色盲模拟.png", "图 8 色盲模拟")
    add_heading(doc, "（五）颜色查询与参考", 2)
    add_paragraph(doc, "“颜色查询”模块用于完成色值换算、参考查找和快速取用，主要操作方式如下：", first_line_indent_cm=None)
    add_bullets(
        doc,
        [
            "颜色格式转换：输入任意常见色值后，可查看对应的 HEX、RGB、HSL、HSV、CMYK 及名称等结果。",
            "高级颜色选择器：通过可视化选色区域调整色相、亮度、透明度等参数，并实时查看计算结果。",
            "渐变色大全：浏览预设渐变方案，复制 CSS 渐变代码或单个色值。",
            "Web 安全色：查看常用 Web Safe Colors 并快速复制色值。",
            "RGB / CMYK 参考表：通过列表或搜索方式查询颜色对照关系，用于设计与印刷参考。",
        ],
    )
    add_image(doc, SCREENSHOT_DIR / "09-颜色格式转换.png", "图 9 颜色格式转换")
    add_heading(doc, "（六）传统色彩", 2)
    add_paragraph(
        doc,
        "“传统色彩”模块提供中国传统色和日本传统色浏览能力。用户可按照分类查看传统色名称、色值与视觉效果，并可点击相应颜色进行复制，用于文化主题设计、包装配色、品牌创意或内容灵感参考。",
    )
    add_heading(doc, "（七）调色板管理", 2)
    add_paragraph(
        doc,
        "“调色板管理”模块用于沉淀和复用个人或项目常用颜色。用户可新建色板、修改色板名称、向色板中添加颜色、删除不需要的颜色，并将当前最近拾取颜色快速加入指定色板。",
    )
    add_paragraph(
        doc,
        "系统支持导入和导出 JSON 格式的色板文件，便于本地备份、项目协作或在不同设备之间迁移。若需要统一归档，也可使用导出全部功能一次性保存全部色板数据。",
    )
    add_heading(doc, "（八）常用快捷键", 2)
    add_paragraph(doc, "桌面端已预置若干全局快捷键，便于快速打开常用模块：", first_line_indent_cm=None)
    add_table(
        doc,
        ["快捷键", "作用"],
        [
            ["Ctrl + Shift + C", "打开屏幕取色页面"],
            ["Ctrl + Shift + R", "打开色盲模拟页面"],
            ["Ctrl + Shift + P", "打开调色板管理页面"],
            ["Ctrl + Shift + Q", "显示或最小化应用窗口"],
        ],
    )
    add_heading(doc, "（九）注意事项", 2)
    add_bullets(
        doc,
        [
            "屏幕取色能力依赖当前运行环境支持情况；若环境不支持相应接口，则该功能可能无法使用。",
            "调色板数据默认保存在本地环境中，导入导出文件是备份和共享的主要方式，并非云端同步。",
            "登录、反馈、检查更新等依赖网络的功能在离线环境下可能受到影响，但大多数本地色彩工具仍可继续使用。",
            "自动更新更适用于正式打包版本；在开发环境中，更新检查通常不会作为主要使用场景。",
            "当图片分辨率较高或文件体积较大时，图片处理速度可能下降，建议根据设备性能选择合适素材。",
            "当前项目以 Windows 桌面端为主要交付方向，其他操作系统的运行方式需另行评估。",
        ],
    )
    add_paragraph(
        doc,
        "本说明书可随项目功能演进持续更新。如后续新增模块、调整菜单结构或变更打包方式，建议同步修订本手册内容，以确保交付文档与实际产品保持一致。",
    )

    return doc


def main() -> None:
    output_path = Path(sys.argv[1]) if len(sys.argv) > 1 else OUTPUT_PATH
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
